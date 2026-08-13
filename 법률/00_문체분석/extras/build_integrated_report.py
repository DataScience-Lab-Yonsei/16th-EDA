from __future__ import annotations

import re
import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
TABLES = ROOT / "results" / "tables"
FIGURES = ROOT / "results" / "figures"
OUTPUT = ROOT / "outputs" / "판결문_문체_통시분석_통합보고서.docx"
GEOMETRY_HELPER = Path(
    r"C:\Users\zzoll\.codex\plugins\cache\openai-primary-runtime\documents"
    r"\26.715.12143\skills\documents\scripts"
)
sys.path.insert(0, str(GEOMETRY_HELPER))
from table_geometry import apply_table_geometry  # noqa: E402


FONT = "Malgun Gothic"
NAVY = "20364F"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2933"
MUTED = "66717D"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"
VERY_LIGHT = "F7F9FB"
WHITE = "FFFFFF"
GOLD = "9A6A00"
RED = "9B1C1C"
TABLE_WIDTH = 9360
TABLE_INDENT = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(
    run,
    size: float | None = None,
    color: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), FONT)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color: str = "D5DAE0", size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = tc_borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            tc_borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:color"), color)


def paragraph_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def paragraph_left_border(paragraph, color: str, size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), size)
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def add_hyperlink(paragraph, text: str, url: str):
    part = paragraph.part
    rel_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), FONT)
    r_fonts.set(qn("w:hAnsi"), FONT)
    r_fonts.set(qn("w:eastAsia"), FONT)
    r_pr.extend([r_fonts, color, underline])
    run.append(r_pr)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def set_image_alt(inline_shape, description: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", description)
    doc_pr.set("title", description)


def create_numbering(doc: Document, kind: str, abstract_id: int, num_id: int) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract_ids = [
        int(element.get(qn("w:abstractNumId")))
        for element in numbering.findall(qn("w:abstractNum"))
        if element.get(qn("w:abstractNumId")) is not None
    ]
    existing_num_ids = [
        int(element.get(qn("w:numId")))
        for element in numbering.findall(qn("w:num"))
        if element.get(qn("w:numId")) is not None
    ]
    abstract_id = max(existing_abstract_ids, default=-1) + 1
    num_id = max(existing_num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if kind == "bullet" else "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "\uf0b7" if kind == "bullet" else "%1.")
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.extend([tabs, indent, spacing])
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    marker_font = "Symbol" if kind == "bullet" else FONT
    r_fonts.set(qn("w:ascii"), marker_font)
    r_fonts.set(qn("w:hAnsi"), marker_font)
    r_fonts.set(qn("w:eastAsia"), marker_font)
    r_pr.append(r_fonts)
    level.extend([start, num_fmt, lvl_text, lvl_jc, p_pr, r_pr])
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def apply_numbering(paragraph, num_id: int) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    p_pr.append(num_pr)


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in tokens.items():
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    caption = doc.styles["Caption"]
    caption.font.name = FONT
    caption._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(MUTED)
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for style_name in ["List Bullet", "List Number"]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = RGBColor.from_string(INK)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def configure_section(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.tab_stops.add_tab_stop(
        Inches(6.5), WD_ALIGN_PARAGRAPH.RIGHT
    )
    r1 = p.add_run("판결문 문체 통시 분석")
    set_run_font(r1, size=8.5, color=MUTED, bold=True)
    r2 = p.add_run("\t내부 검토용 통합 보고서")
    set_run_font(r2, size=8.5, color=MUTED)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    add_page_number(fp)


def add_title_block(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(72)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("분석 보고서")
    set_run_font(run, size=11, color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("판결문의 말은\n어떻게 변해왔는가")
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run("형사 판결문 이유부의 문체 변화, 1980–2025")
    set_run_font(run, size=14, color=DARK_BLUE)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(28)
    run = p.add_run("명사·동사·조사 비율 · 문장 길이 · 표현 변이 · 견고성 검증")
    set_run_font(run, size=10.5, color=MUTED)

    cover_image = doc.add_picture(
        str(FIGURES / "F4_metric_heatmap.png"), width=Inches(5.7)
    )
    set_image_alt(cover_image, "판결문 문체 지표의 연도별 변화 히트맵")
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.paragraphs[-1].paragraph_format.space_after = Pt(24)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("작성일  2026년 7월 23일")
    set_run_font(run, size=10.5, color=MUTED, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("분석 대상  공개 형사 판결문 이유부 18,778건")
    set_run_font(run, size=9.5, color=MUTED)
    doc.add_page_break()


def add_bullet(doc: Document, text: str, num_id: int, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2)
    else:
        set_run_font(p.add_run(text))
    return p


def add_numbered(doc: Document, text: str, num_id: int):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.167
    set_run_font(p.add_run(text))
    return p


def add_callout(doc: Document, label: str, text: str, color: str = BLUE):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.15
    paragraph_shading(p, LIGHT_BLUE)
    paragraph_left_border(p, color)
    r1 = p.add_run(f"{label}  ")
    set_run_font(r1, bold=True, color=color)
    r2 = p.add_run(text)
    set_run_font(r2, color=INK)
    return p


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[int],
    alignments: list[str] | None = None,
    font_size: float = 9.3,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.allow_autofit = False
    header = table.rows[0]
    set_repeat_table_header(header)
    for idx, text in enumerate(headers):
        cell = header.cells[idx]
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_shading(cell, LIGHT_GRAY)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=font_size, bold=True, color=NAVY)

    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            alignment = (alignments or ["left"] * len(headers))[idx]
            p.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }[alignment]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            run = p.add_run(str(value))
            set_run_font(run, size=font_size)

    apply_table_geometry(
        table,
        widths,
        table_width_dxa=TABLE_WIDTH,
        indent_dxa=TABLE_INDENT,
        cell_margins_dxa=CELL_MARGINS,
    )
    after = doc.add_paragraph()
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.space_after = Pt(2)
    return table


def add_figure(doc: Document, filename: str, caption: str, width: float = 6.25):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    image = p.add_run().add_picture(str(FIGURES / filename), width=Inches(width))
    set_image_alt(image, caption)
    cap = doc.add_paragraph(caption, style="Caption")
    cap.paragraph_format.keep_with_next = False


def add_source_line(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=8.5, color=MUTED)


def percent(value: float, digits: int = 2) -> str:
    return f"{value * 100:.{digits}f}%"


def pp(value: float, digits: int = 2) -> str:
    return f"{value * 100:+.{digits}f}%p"


def extract_example(markdown: str, year: int, max_chars: int = 560) -> str:
    match = re.search(
        rf"## {year}년[\s\S]*?> ([^\r\n]+)",
        markdown,
    )
    if not match:
        return ""
    text = re.sub(r"\s+", " ", match.group(1)).strip()
    return text if len(text) <= max_chars else text[: max_chars - 1].rstrip() + "…"


def add_quote(doc: Document, label: str, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.right_indent = Inches(0.12)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.2
    paragraph_shading(p, VERY_LIGHT)
    paragraph_left_border(p, "AAB7C4", size="12")
    r1 = p.add_run(label + "\n")
    set_run_font(r1, size=10.5, color=DARK_BLUE, bold=True)
    r2 = p.add_run(text)
    set_run_font(r2, size=9.7, color=INK)


def build_report() -> Path:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    bullet_id = 0
    decimal_id = 0

    standardized = pd.read_csv(TABLES / "annual_standardized_metrics.csv")
    trends = pd.read_csv(TABLES / "trend_models.csv")
    variants = pd.read_csv(TABLES / "variant_summary.csv")
    subgroup = pd.read_csv(TABLES / "robustness_subgroups.csv")
    balanced = pd.read_csv(TABLES / "robustness_balanced_resampling.csv")
    year_fe = pd.read_csv(TABLES / "year_fe_comparison.csv")
    placebo = pd.read_csv(TABLES / "placebo_tests.csv")
    breakpoints = pd.read_csv(TABLES / "breakpoint_bootstrap.csv")
    qc = pd.read_csv(TABLES / "qc_1980s_summary.csv")
    variant_qc = pd.read_csv(TABLES / "variant_context_summary.csv")
    hac = pd.read_csv(TABLES / "hac_annual_regression.csv")
    aggregation = pd.read_csv(TABLES / "aggregation_sensitivity.csv")
    length_comparison = pd.read_csv(TABLES / "length_control_comparison.csv")
    examples_md = (ROOT / "results" / "examples" / "matched_examples.md").read_text(
        encoding="utf-8"
    )

    add_title_block(doc)

    doc.add_heading("요약", level=1)
    add_callout(
        doc,
        "한 문장 결론",
        "1980–2025년 공개 형사 판결문 이유부는 문장이 약 절반으로 짧아지는 동시에, "
        "동사·서술어보다 명사를 더 많이 쓰는 방향으로 이동했다.",
    )
    add_bullet(
        doc,
        "문장 길이: 구성 보정 기준 문장당 63.43어절(1985)에서 30.51어절(2025)로 51.9% 감소했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "어휘·문법 구성: 명사 비율은 40.78%에서 43.88%로 3.10%p 상승했고, "
        "동사 비율은 5.41%에서 4.51%로 0.90%p 하락했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "조사 비율: 19.69%에서 19.05%로 0.64%p 낮아졌다. 변화 폭은 문장 길이나 명사성보다 작다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "견고성: 여섯 범죄군 모두 핵심 방향이 같았고, 균형 재표집·연도 더미 회귀·Newey-West 회귀선에서도 방향이 유지됐다. "
        "다만 2심의 명사 비율은 뚜렷한 증가가 아니라 거의 평평했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "해석 한계: 이 결과는 공개된 판결문 코퍼스의 변화다. 공개 대상, 수집 방식, 구두점·입력 관행의 변화와 완전히 분리할 수 없다.",
        bullet_id,
    )
    add_callout(
        doc,
        "데이터 확인",
        "팀장 전달 내용과 달리 제공된 crime.zip에는 1947–1949년 판결문이 없다. "
        "현재 파일의 최초 연도는 1950년이며, 본 분석은 표본 구성이 비교적 안정적인 1980–2025년을 사용했다.",
        color=GOLD,
    )

    doc.add_page_break()

    doc.add_heading("1. 분석 질문과 표본", level=1)
    p = doc.add_paragraph(
        "팀장 요청의 핵심은 “판결문의 말이 시대별로 어떻게 변했는가”를 연도별 추세로 보여주는 것이다. "
        "이에 따라 명사·동사·조사 비율과 문장 길이를 중심 지표로 두고, 표현 변이쌍을 별도 트랙으로 분석했다."
    )
    p.paragraph_format.keep_with_next = True

    doc.add_heading("1.1 분석 표본 정의", level=2)
    sample_rows = [
        ["원자료 형사 판결문", "20,984건", "제공 ZIP에 포함된 판결문"],
        ["이유부 추출 성공", "20,919건", "전 구간 약 99% 이상"],
        ["형태소 150개 이상", "20,743건", "지나치게 짧은 문서 제외"],
        ["본 분석", "18,778건", "1980–2025년 이유부"],
        ["보조 탐색", "1950–1979년", "초기 표본이 작아 결론용으로 사용하지 않음"],
        ["제외", "2026년 49건", "부분 연도"],
    ]
    add_table(
        doc,
        ["단계", "규모", "적용 기준"],
        sample_rows,
        [2550, 1700, 5110],
        ["left", "center", "left"],
    )
    add_source_line(
        doc,
        "주: 제공된 C:\\Users\\zzoll\\Downloads\\crime.zip과 분석용 data/raw/crime.zip의 SHA-256은 동일하다. "
        "두 파일 모두 1947–1949년 자료를 포함하지 않는다.",
    )

    doc.add_heading("1.2 왜 원시 시계열만 보면 안 되는가", level=2)
    p = doc.add_paragraph(
        "연도별 표본 구성은 고정되어 있지 않다. 특히 심급과 범죄군의 비중, 판결문 길이와 공개되는 문서의 성격이 "
        "시대별로 달라진다. 따라서 원시 평균의 변화에는 문체 변화와 표본 구성 변화가 함께 들어간다."
    )
    add_figure(
        doc,
        "F1_composition.png",
        "그림 1. 연도별 표본 수, 3심 비중, 문서 길이와 범죄군 구성",
        width=6.2,
    )
    add_callout(
        doc,
        "분석 원칙",
        "원시 추세는 진단용으로 보고, 결론은 심급×범죄군 구성을 고정한 직접표준화와 통제 회귀를 함께 사용한다.",
    )
    doc.add_page_break()

    doc.add_heading("2. 분석 방법", level=1)
    method_rows = [
        ["관측 단위", "판결문 1건"],
        ["시간축", "교정된 선고연도, 1980–2025"],
        ["본문 범위", "주문을 제외한 이유부"],
        ["핵심 지표", "명사·동사·서술어·조사 비율, 명사성, 문장당 어절·글자 수"],
        ["직접표준화", "전 기간 심급×범죄군 고정 가중치"],
        ["회귀 통제", "심급, 범죄군, 문서 성격, 법원 고정효과, 로그 문서 길이"],
        ["불확실성", "법원 군집 표준오차, BH 다중검정 보정"],
        ["연도 계열 자기상관", "구성 보정 연도값 회귀에 Newey-West HAC(lag 3) 적용"],
        ["추가 검증", "연도 더미 회귀, 범죄군·심급 분할, 균형 재표집 500회, 위약검정"],
        ["변화점", "3년 블록 잔차 부트스트랩 1,000회"],
    ]
    add_table(
        doc,
        ["항목", "적용 방식"],
        method_rows,
        [2200, 7160],
        ["left", "left"],
        font_size=9.4,
    )
    doc.add_heading("2.1 지표 해석", level=2)
    add_bullet(
        doc,
        "명사 비율은 전체 분석 형태소 중 명사가 차지하는 비중이다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "명사성은 명사/(명사+서술어)다. 값이 높을수록 사건과 행위를 동사로 풀기보다 명사로 묶어 표현하는 경향이 강하다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "문장당 어절 수는 문장 경계 하나 안에 들어가는 띄어쓰기 단위의 평균 개수다. "
        "오래된 판결문의 구두점 관행에도 영향을 받을 수 있다.",
        bullet_id,
    )
    add_callout(
        doc,
        "중요",
        "‘문장이 짧아졌다’와 ‘글이 쉬워졌다’는 같은 말이 아니다. 짧은 문장과 높은 명사성이 동시에 나타날 수 있다.",
        color=GOLD,
    )

    doc.add_heading("3. 핵심 시계열 결과", level=1)
    add_figure(
        doc,
        "F2_main_trends.png",
        "그림 2. 원자료 평균과 심급×범죄군 구성 보정 추세",
        width=6.2,
    )

    values = (
        standardized.loc[standardized["year"].isin([1985, 2025])]
        .pivot(index="metric", columns="year", values="estimate")
    )
    comparison_rows = []
    labels = {
        "noun_share": "명사 비율",
        "verb_share": "동사 비율",
        "predicate_share": "서술어 비율",
        "particle_share": "조사 비율",
        "nominality": "명사성",
        "sent_eojeol_mean": "문장당 어절 수",
        "sent_char_mean": "문장당 글자 수",
    }
    for metric, label in labels.items():
        old = float(values.loc[metric, 1985])
        new = float(values.loc[metric, 2025])
        if metric.endswith("share") or metric == "nominality":
            change = f"{(new - old) * 100:+.2f}%p"
            old_text, new_text = percent(old), percent(new)
        else:
            change = f"{(new / old - 1) * 100:+.1f}%"
            old_text, new_text = f"{old:.2f}", f"{new:.2f}"
        comparison_rows.append([label, old_text, new_text, change])
    add_table(
        doc,
        ["지표", "1985", "2025", "변화"],
        comparison_rows,
        [3300, 1900, 1900, 2260],
        ["left", "right", "right", "right"],
    )

    doc.add_heading("3.1 핵심 해석", level=2)
    add_bullet(
        doc,
        "명사 중심성 증가: 명사 비율과 명사성이 상승하고, 동사·서술어 비율은 하락했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "문장 경계 증가: 문장당 어절과 글자 수가 약 절반으로 줄었다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "조사 비율은 완만하게 하락했다. 통계적으로는 유의하지만 효과크기는 문장 길이보다 작다.",
        bullet_id,
    )
    add_callout(
        doc,
        "해석",
        "최근 판결문은 한 문장을 더 잘게 끊지만, 각 문장 안에서는 사건 요소와 법률 개념을 명사로 묶는 경향이 더 강하다.",
    )
    doc.add_heading("3.2 보정 회귀선과 10년당 변화", level=2)
    add_figure(
        doc,
        "F10_adjusted_regression_lines.png",
        "그림 3. 구성 보정 연도값, 선형 회귀선과 Newey-West 95% 신뢰구간",
        width=6.0,
    )
    hac_key = hac.set_index("metric")
    add_callout(
        doc,
        "회귀선 결과",
        "연도별 구성 보정값에 직선을 적합하고 3년 자기상관을 허용했을 때 10년당 변화는 "
        f"명사 {hac_key.loc['noun_share', 'effect_per_decade'] * 100:+.3f}%p, "
        f"동사 {hac_key.loc['verb_share', 'effect_per_decade'] * 100:+.3f}%p, "
        f"조사 {hac_key.loc['particle_share', 'effect_per_decade'] * 100:+.3f}%p, "
        f"문장 길이 {hac_key.loc['sent_eojeol_mean', 'effect_per_decade']:+.2f}어절이었다. "
        "네 지표 모두 BH 보정 후 유의했다.",
    )
    doc.add_page_break()
    trend_key = trends.loc[
        trends["length_control"].astype(str).str.lower().eq("true")
        & trends["metric"].isin(["noun_share", "verb_share", "particle_share", "sent_eojeol_mean"])
    ].set_index("metric")
    trend_rows = []
    for metric in ["noun_share", "verb_share", "particle_share", "sent_eojeol_mean"]:
        row = trend_key.loc[metric]
        if metric.endswith("share"):
            effect = pp(float(row["effect_per_decade"]), 3)
            ci = f"{pp(float(row['ci_low']), 3)} ~ {pp(float(row['ci_high']), 3)}"
        else:
            effect = f"{float(row['effect_per_decade']):+.2f}어절"
            ci = f"{float(row['ci_low']):+.2f} ~ {float(row['ci_high']):+.2f}"
        trend_rows.append(
            [
                row["metric_label"],
                effect,
                ci,
                f"{float(row['q_value']):.2g}",
            ]
        )
    add_table(
        doc,
        ["지표", "문서 회귀 10년당", "95% 신뢰구간", "BH q값"],
        trend_rows,
        [2800, 1900, 2900, 1760],
        ["left", "right", "center", "right"],
    )
    p = doc.add_paragraph(
        "회귀선은 구성 보정된 연도값의 장기 방향을 보여주고, 표는 개별 판결문 수준에서 문서 길이·심급·범죄군·"
        "문서 성격·법원 고정효과를 통제한 결과다. 두 접근의 방향이 일치하며 BH 다중검정 보정 후 유의했다."
    )

    add_figure(
        doc,
        "F6_effect_forest.png",
        "그림 4. 전체 지표의 보정 회귀 10년당 효과크기와 95% 신뢰구간",
        width=5.9,
    )
    p = doc.add_paragraph(
        "문장 길이와 서술어·명사성의 변화가 표준편차 단위에서 가장 크고, 조사와 한자 비율의 효과크기는 상대적으로 작다."
    )

    doc.add_heading("3.3 전체 지표의 변화 모양", level=2)
    add_figure(
        doc,
        "F3_metric_small_multiples.png",
        "그림 5. 구성 보정 문체 지표의 공통 SD 척도 비교",
        width=6.0,
    )
    add_figure(
        doc,
        "F4_metric_heatmap.png",
        "그림 6. 지표별 변화 시점 히트맵",
        width=6.2,
    )
    p = doc.add_paragraph(
        "명사성·동사·서술어 변화는 1990년대 이후 뚜렷해지고, 문장 길이는 2000년대 중반 이후 급격히 낮아진다. "
        "모든 지표가 같은 시점에 움직인 것은 아니다."
    )

    doc.add_heading("4. 표현 변이", level=1)
    add_figure(
        doc,
        "F5_variant_transitions.png",
        "그림 7. 여섯 표현 변이쌍의 기존형 비율",
        width=6.15,
    )
    variant_rows = []
    for row in variants.itertuples(index=False):
        t50 = (
            f"{row.t50:.0f}년"
            if pd.notna(row.t50) and 1980 <= row.t50 <= 2025
            else "관측창 밖/미추정"
        )
        variant_rows.append(
            [
                f"{row.label_old} → {row.label_new}",
                f"{int(row.total_old):,}",
                f"{int(row.total_new):,}",
                t50,
            ]
        )
    add_table(
        doc,
        ["변이쌍", "기존형", "새 형태", "50% 교체 시점"],
        variant_rows,
        [3500, 1700, 1700, 2460],
        ["left", "right", "right", "center"],
        font_size=9.1,
    )
    add_callout(
        doc,
        "가장 분명한 교체",
        "‘아니하- → 않-’은 약 1991년, ‘하지 아니하- → 하지 않-’은 약 1995년에 기존형 비율이 50% 아래로 내려간 것으로 추정된다.",
    )

    new_quote = variant_qc.loc[variant_qc["form"].eq("new")].set_index("key")
    p = doc.add_paragraph(
        "다른 축약형은 해석에 주의가 필요하다. 문맥 표본 12건씩을 점검했을 때 새 형태 ‘됐-’의 58%, "
        "‘돼’의 67%, ‘했-’과 ‘해야’의 각 25%가 인용문 안에 있었다. "
        "따라서 이 네 쌍은 판결문 서술 자체보다 피고인·증인 발화의 인용 비중 변화가 일부 섞였을 수 있다. "
        "반면 ‘하지 않-’ 표본에서는 인용문 비중이 0%로 나타나 부정 표현 교체가 상대적으로 더 안정적인 신호다."
    )
    add_callout(
        doc,
        "중복 주의",
        "‘아니하- → 않-’의 약 절반은 더 긴 ‘하지 아니하- → 하지 않-’ 구성 안에 포함된다. "
        "두 결과를 독립적인 변화 두 개로 합산하면 안 된다.",
        color=GOLD,
    )
    doc.add_heading("5. 실제 판결문 사례", level=1)
    p = doc.add_paragraph(
        "같은 3심·재산범죄 안에서 각 연도의 명사 비율과 문장 길이가 중앙값에 가까운 판결문을 골랐다. "
        "사례는 방향을 이해하기 위한 예시이며 전체 판결문을 대표하지 않는다."
    )
    add_table(
        doc,
        ["연도", "명사", "동사", "조사", "문장당 어절"],
        [
            ["1985", "41.1%", "4.4%", "18.2%", "49.0"],
            ["2025", "44.3%", "4.1%", "18.5%", "27.5"],
        ],
        [1600, 1800, 1800, 1800, 2360],
        ["center", "right", "right", "right", "right"],
    )
    add_quote(doc, "1985년 예시", extract_example(examples_md, 1985))
    add_quote(doc, "2025년 예시", extract_example(examples_md, 2025))
    p = doc.add_paragraph(
        "두 사례만으로 시대 차이를 확정할 수는 없지만, 2025년 예시는 항목 번호와 짧은 문장 경계가 더 자주 나타나고 "
        "법률 개념은 여전히 명사 중심으로 배열된다. 통계 결과는 이러한 인상이 전체 코퍼스에서도 장기적으로 나타나는지를 검증한다."
    )

    doc.add_heading("6. 견고성 검증", level=1)
    doc.add_heading("6.1 보정 방법을 바꿔도 같은가", level=2)
    add_figure(
        doc,
        "F8_adjustment_robustness.png",
        "그림 8. 직접표준화, 연도 더미 회귀, 균형 재표집 비교",
        width=5.85,
    )
    corr_text = ", ".join(
        f"{row.metric_label} r={row.correlation_with_direct_standardization:.3f}"
        for row in year_fe.itertuples(index=False)
    )
    p = doc.add_paragraph(
        "직접표준화와 연도 더미 회귀의 연도별 곡선 상관은 "
        + corr_text
        + "였다. 균형 재표집도 명사 증가, 동사·조사·문장 길이 감소를 재현했다."
    )
    balanced_key = balanced.set_index("metric")
    balanced_rows = []
    for metric in ["noun_share", "verb_share", "particle_share", "sent_eojeol_mean"]:
        row = balanced_key.loc[metric]
        if metric.endswith("share"):
            est = pp(float(row["effect_per_decade"]), 3)
            ci = f"{pp(float(row['ci_low']), 3)} ~ {pp(float(row['ci_high']), 3)}"
        else:
            est = f"{float(row['effect_per_decade']):+.2f}어절"
            ci = f"{float(row['ci_low']):+.2f} ~ {float(row['ci_high']):+.2f}"
        balanced_rows.append([row["metric_label"], est, ci])
    add_table(
        doc,
        ["지표", "균형 재표집 10년당 변화", "부트스트랩 95% 구간"],
        balanced_rows,
        [3000, 2800, 3560],
        ["left", "right", "center"],
    )
    add_source_line(
        doc,
        "균형 재표집: 1985–2025년에 모두 관측된 8개 심급×범죄군 셀에서 셀·연도당 20건을 복원추출, 500회 반복.",
    )

    doc.add_heading("6.2 범죄군·심급으로 나눠도 같은가", level=2)
    add_figure(
        doc,
        "F7_subgroup_robustness.png",
        "그림 9. 범죄군·심급별 10년당 보정 변화량과 95% 신뢰구간",
        width=6.15,
    )
    add_bullet(
        doc,
        "여섯 범죄군 모두 명사 비율은 증가하고, 동사·조사 비율과 문장 길이는 감소했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "1심과 3심의 명사 비율은 증가했지만, 2심은 -0.009%p/10년으로 거의 평평했고 신뢰구간이 0을 포함했다.",
        bullet_id,
    )
    add_bullet(
        doc,
        "모든 심급에서 동사·조사 비율과 문장 길이 감소 방향은 유지됐다.",
        bullet_id,
    )
    add_source_line(
        doc,
        "주: 3심 표본은 사실상 대법원 한 곳이므로 법원 군집 표준오차 대신 HC3 이분산 강건 표준오차를 사용했다.",
    )

    doc.add_heading("6.3 위약검정", level=2)
    placebo_row = placebo.iloc[0]
    p = doc.add_paragraph(
        "판례 ID의 SHA-256 해시값을 문체와 무관한 위약 결과변수로 두고 같은 회귀를 적용했다. "
        f"10년당 효과는 {float(placebo_row.effect_per_decade):+.4f}, "
        f"95% 신뢰구간 {float(placebo_row.ci_low):+.4f}~{float(placebo_row.ci_high):+.4f}, "
        f"p={float(placebo_row.p_value):.3f}로 추세가 검출되지 않았다."
    )
    add_callout(
        doc,
        "견고성 결론",
        "표본 구성을 통제하는 방법과 분석 하위집단을 바꿔도 핵심 방향은 대체로 유지된다. "
        "다만 2심 명사 비율과 변이쌍의 인용문 문제는 결론의 범위를 제한한다.",
    )
    doc.add_heading("7. 변화점과 제도적 맥락", level=1)
    add_figure(
        doc,
        "F9_breakpoint_uncertainty.png",
        "그림 10. 탐색적 변화점 후보와 3년 블록 부트스트랩 95% 구간",
        width=6.15,
    )
    p = doc.add_paragraph(
        "문장당 어절·글자 수의 최적 변화점 후보는 2004년이며, 부트스트랩 95% 구간은 대략 2000–2008년이다. "
        "명사성·동사·서술어는 1990년대 초반에 변화점 후보가 나타나지만, 명사·조사 비율의 구간은 2017년까지 넓어 "
        "특정 연도를 단정하기 어렵다."
    )
    timeline_rows = [
        [
            "2001",
            "신임 법관 연수에서 ‘새로운 판결서 작성방식’을 교육한 공식 기록이 확인된다.",
            "대법원 법원동정",
        ],
        [
            "2004",
            "특허법원 판사들이 난해한 설시를 지양하고 명확한 설명과 일본식·한자식 용어 순화를 논의했다.",
            "대법원 법원동정",
        ],
        [
            "2013·2015",
            "형사 및 민사 등 판결서 인터넷 공개 범위가 단계적으로 확대돼 공개 표본 환경이 달라졌다.",
            "대법원 보도자료",
        ],
        [
            "2022",
            "사법행정자문회의에서 국민이 이해하기 쉬운 판결서 작성 방식과 개선 방안을 연구·검토할 필요가 논의됐다.",
            "대법원 회의결과",
        ],
    ]
    add_table(
        doc,
        ["시점", "공식 기록에서 확인되는 내용", "출처"],
        timeline_rows,
        [1250, 6200, 1910],
        ["center", "left", "center"],
        font_size=8.8,
    )
    add_callout(
        doc,
        "인과 해석 금지",
        "2004년 전후 내부 교육·논의와 문장 길이 변화점의 시기가 겹치지만, 이것만으로 제도 변화가 문체 변화를 일으켰다고 볼 수 없다. "
        "같은 시기에 전산화·공개·편집 관행도 함께 변했다.",
        color=GOLD,
    )

    doc.add_heading("8. 품질 점검", level=1)
    qc_lookup = qc.set_index("item")["value"]
    qc_rows = [
        ["층화 표본", f"{int(qc_lookup['sample_size'])}건", "1980–1989년 연도별 20건"],
        [
            "재계산 일치",
            f"{int(qc_lookup['count_recalculation_mismatches'])}건 불일치",
            "형태소·문장·명사·동사·조사 수 재현",
        ],
        [
            "문장 경계 위험",
            f"{int(qc_lookup['boundary_risk_documents'])}건",
            "문장당 100어절 초과 또는 긴 문서의 경계 부족",
        ],
        [
            "한자 고비율",
            f"{int(qc_lookup['high_hanja_documents'])}건",
            "표본 중 한자 비율 1% 초과 없음",
        ],
    ]
    add_table(
        doc,
        ["점검 항목", "결과", "판정 기준"],
        qc_rows,
        [2200, 2200, 4960],
        ["left", "center", "left"],
    )
    p = doc.add_paragraph(
        "재계산 불일치는 없었으나 200건 중 15건은 문장당 100어절을 넘었다. 위험 표본을 직접 읽어보면 내용 손상보다는 "
        "쉼표·열거를 길게 이어 쓰는 당시 판결문 관행이 주된 원인이었다. 따라서 문장 길이 감소는 실제 문체 변화와 "
        "구두점·편집 관행 변화를 함께 반영한 결과로 해석하는 것이 안전하다."
    )

    doc.add_heading("9. 한계", level=1)
    limitations = [
        "공개 판결문 표본의 변화이며 사법부 전체 판결문을 직접 대표하지 않는다.",
        "현재 제공 파일에는 1947–1949년이 없고, 1950–1979년은 연도별 표본이 작아 보조 탐색에만 사용했다.",
        "제공 파일이 형사 판결문으로 한정되어 민사·조세 단독 견고성 검증은 수행할 수 없었다. 해당 검증에는 전체 판례 마트가 필요하다.",
        "심급·범죄군·법원·문서 길이를 통제했지만 관측되지 않은 사건 난이도나 작성 관행 차이는 남는다.",
        "형태소 분석기는 현대 한국어 모델이므로 역사적 표기 분석에는 한계가 있다.",
        "문장 길이는 구두점·전산화·입력 관행 변화에도 민감하다.",
        "표현 변이 중 축약형은 인용문에 자주 등장하며, ‘아니하/않’ 두 지표는 약 절반이 중첩된다.",
        "변화점은 탐색적 요약이고 정책 효과 추정치가 아니다.",
    ]
    for item in limitations:
        add_bullet(doc, item, bullet_id)

    doc.add_heading("10. 결론", level=1)
    add_callout(
        doc,
        "최종 결론",
        "공개 형사 판결문 이유부는 장기적으로 더 짧게 끊어 쓰이지만, 문장 내부의 표현은 더 명사 중심적으로 이동했다. "
        "이 방향은 사건 구성 보정, 문서 단위 통제 회귀, Newey-West 회귀선, 범죄군 분할과 균형 재표집에서 대체로 유지됐다.",
    )
    p = doc.add_paragraph(
        "따라서 보고 시에는 “판결문이 단순히 쉬워졌다”라고 표현하기보다, "
        "“문장 경계는 짧아졌고 어휘·문법 구성은 명사 중심으로 이동했다”라고 정리하는 것이 가장 정확하다."
    )
    doc.add_heading("팀장 보고용 권장 문장", level=2)
    add_quote(
        doc,
        "보고 문안",
        "1980–2025년 공개 형사 판결문 이유부를 분석한 결과, 심급과 범죄군 구성을 보정한 뒤에도 "
        "문장당 어절 수는 약 절반으로 감소했고 명사 비율은 약 3.1%p 증가했다. "
        "즉 최근 판결문은 문장을 더 짧게 끊지만, 표현은 동사보다 명사 중심으로 구성되는 경향이 강하다. "
        "다만 이는 공개 판결문 코퍼스의 변화이며 사법부 전체의 인과적 변화로 해석할 수는 없다.",
    )

    doc.add_page_break()
    doc.add_heading("부록 A. 지표 정의", level=1)
    metric_rows = [
        ["명사 비율", "명사 수 / 분석 형태소 수", "개념·대상 중심 구성"],
        ["동사 비율", "일반동사(VV) 수 / 분석 형태소 수", "행위 서술 비중"],
        ["서술어 비율", "동사·형용사·보조용언·지정사 / 분석 형태소 수", "전체 서술 기능"],
        ["조사 비율", "조사 수 / 분석 형태소 수", "문법 관계 표지"],
        ["명사성", "명사 / (명사 + 서술어)", "명사 중심성"],
        ["문장당 어절", "어절 수 / 문장 수", "문장 경계의 길이"],
        ["문장당 글자", "공백 제외 글자 수 / 문장 수", "문장 물리적 길이"],
    ]
    add_table(
        doc,
        ["지표", "계산식", "해석"],
        metric_rows,
        [2300, 4000, 3060],
        ["left", "left", "left"],
        font_size=9.0,
    )

    doc.add_heading("부록 B. 원자료·집계 민감도", level=1)
    add_figure(
        doc,
        "F11_raw_median_iqr.png",
        "그림 11. 원자료 중앙값·IQR과 문서평균",
        width=6.0,
    )
    p = doc.add_paragraph(
        "왜도가 있는 원자료에서도 중앙값과 평균의 장기 방향은 대체로 일치한다. "
        "다만 문장 길이는 연도 내 분산이 크므로 원시 곡선만으로 결론을 내리지 않았다."
    )
    add_figure(
        doc,
        "F12_aggregation_sensitivity.png",
        "그림 12. 문서평균과 토큰·문장 풀링 집계의 원자료 추세",
        width=6.0,
    )
    aggregation_rows = []
    for row in aggregation.itertuples(index=False):
        if str(row.metric).endswith("share"):
            doc_effect = f"{float(row.document_mean_effect_per_decade) * 100:+.3f}%p"
            pooled_effect = f"{float(row.pooled_effect_per_decade) * 100:+.3f}%p"
        else:
            doc_effect = f"{float(row.document_mean_effect_per_decade):+.2f}어절"
            pooled_effect = f"{float(row.pooled_effect_per_decade):+.2f}어절"
        aggregation_rows.append(
            [
                row.metric_label,
                doc_effect,
                pooled_effect,
                "동일" if bool(row.same_slope_direction) else "상이",
            ]
        )
    add_table(
        doc,
        ["지표", "문서평균 10년당", "풀링 10년당", "방향"],
        aggregation_rows,
        [2600, 2500, 2500, 1760],
        ["left", "right", "right", "center"],
    )
    add_source_line(
        doc,
        "주: 풀링 비교는 긴 문서가 더 큰 가중치를 갖는 원자료 진단이다. 결론에는 문서평균과 구성 보정값을 사용했다.",
    )

    doc.add_page_break()
    doc.add_heading("부록 C. 길이 통제 민감도", level=1)
    length_rows = []
    for row in length_comparison.itertuples(index=False):
        length_rows.append(
            [
                row.metric_label,
                f"{float(row.without_length_effect_sd_per_decade):+.3f}",
                f"{float(row.with_length_effect_sd_per_decade):+.3f}",
                "동일" if bool(row.same_direction) else "상이",
            ]
        )
    add_table(
        doc,
        ["지표", "길이 통제 없음", "길이 통제 있음", "방향"],
        length_rows,
        [2600, 2300, 2300, 2160],
        ["left", "right", "right", "center"],
        font_size=9.1,
    )
    p = doc.add_paragraph(
        "표의 수치는 10년당 문서 간 표준편차(SD) 변화량이다. 모든 지표가 문서 길이 통제 여부와 관계없이 같은 방향을 보였다."
    )

    doc.add_heading("부록 D. 재현 순서", level=1)
    for text in [
        "01_prepare_data.py: 날짜 교정, 이유부 추출, 심급·범죄군 규칙 적용",
        "02_extract_metrics.py: Kiwi 형태소 분석과 문체 지표 계산",
        "03_analyze_timeseries.py: 원시·표준화 추세, 보정 회귀, 표현 변이",
        "04_create_figures.py 및 05_extract_examples.py: 기본 그림과 동일 조건 사례",
        "06_robustness.py: 하위집단, 연도 더미, 균형 재표집, 위약·변화점 검증",
        "07_create_robustness_figures.py 및 08_qc_context.py: 검증 그림과 문맥 점검",
        "09_supplementary_analysis.py: 회귀선·HAC, 중앙값·IQR, 풀링 및 길이 통제 민감도",
    ]:
        add_numbered(doc, text, decimal_id)

    doc.add_page_break()
    doc.add_heading("부록 E. 공식 자료", level=1)
    sources = [
        (
            "대법원 법원동정, 2001.6.1. — 신임 법관 연수 ‘새로운 판결서 작성방식’",
            "https://scourt.go.kr/portal/news/NewsViewAction.work?gubun=5&searchOption=&searchWord=&seqnum=62",
        ),
        (
            "대법원 법원동정, 2004.4. — 명확한 설시와 일본식·한자식 용어 순화 논의",
            "https://www.scourt.go.kr/portal/news/NewsViewAction.work?gubun=5&searchOption=&searchWord=&seqnum=113",
        ),
        (
            "대법원 보도자료, 2018.8.21. — 판결서 공개범위 확대 추진 경과",
            "https://scourt.go.kr/portal/news/NewsViewAction.work?gubun=6&searchOption=&searchWord=&seqnum=1550",
        ),
        (
            "대법원 사법행정자문회의 회의결과 — 판결서 작성 개선 방안 논의",
            "https://www.scourt.go.kr/supreme/news/NewsViewAction2.work?gubun=944&searchOption=&searchWord=&seqnum=45",
        ),
    ]
    for label, url in sources:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        add_hyperlink(p, label, url)

    doc.core_properties.title = "판결문의 말은 어떻게 변해왔는가"
    doc.core_properties.subject = "형사 판결문 이유부 문체 통시 분석, 1980–2025"
    doc.core_properties.author = "판결문 문체 통시 분석 프로젝트"
    doc.core_properties.keywords = "판결문, 통시 분석, 형태소, 명사 비율, 문장 길이"
    doc.core_properties.comments = "통합 분석 보고서"

    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build_report()
    print(path)
